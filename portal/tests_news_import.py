import base64
import io
import shutil
import tempfile
import zipfile
from pathlib import Path

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document

from news.models import Category, News

from .models import PortalAuditLog


class PortalNewsImportTests(TestCase):
    TINY_PNG = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9sotWJkAAAAASUVORK5CYII="
    )

    def setUp(self):
        super().setUp()
        self.temp_media = tempfile.mkdtemp()
        self.addCleanup(shutil.rmtree, self.temp_media, ignore_errors=True)
        self.override_media = override_settings(MEDIA_ROOT=self.temp_media)
        self.override_media.enable()
        self.addCleanup(self.override_media.disable)

        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="portal_import_admin",
            password="password123",
            is_staff=True,
            is_superuser=True,
        )
        self.client.force_login(self.user)
        self.category = Category.objects.create(name="Imported")

    def _make_docx_upload(self, title, body_lines):
        document = Document()
        document.add_paragraph(title)
        for line in body_lines:
            document.add_paragraph(line)

        image_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        image_path = Path(image_file.name)
        image_file.close()
        self.addCleanup(lambda: image_path.unlink(missing_ok=True))
        image_path.write_bytes(self.TINY_PNG)
        document.add_picture(str(image_path))

        payload = io.BytesIO()
        document.save(payload)
        payload.seek(0)
        return SimpleUploadedFile(
            "portal-import.docx",
            payload.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def _make_zip_upload(self):
        payload = io.BytesIO()
        with zipfile.ZipFile(payload, "w") as archive:
            archive.writestr("gallery/extra.png", self.TINY_PNG)
        payload.seek(0)
        return SimpleUploadedFile("extra-images.zip", payload.getvalue(), content_type="application/zip")

    def test_news_import_creates_news_and_audit_log(self):
        response = self.client.post(
            reverse("portal:news_import"),
            data={
                "source_file": self._make_docx_upload(
                    "BÀI TEST IMPORT",
                    [
                        "Đây là đoạn mở đầu.",
                        "Danh sách học sinh đạt giải:",
                        "Nguyễn Văn A – Giải Nhất",
                    ],
                ),
                "extra_images_zip": self._make_zip_upload(),
                "category": self.category.pk,
                "is_featured": "on",
                "overwrite_existing": "on",
            },
            HTTP_HOST="localhost",
        )

        self.assertEqual(response.status_code, 302)
        news = News.objects.get(slug="bai-test-import")
        self.assertEqual(news.category, self.category)
        self.assertTrue(news.is_featured)
        self.assertIn("<figure><img", news.content)
        self.assertIn("<ul>", news.content)
        self.assertTrue(news.thumbnail.name.startswith("news/imported/"))
        self.assertTrue(news.source_document.name.startswith("news/source-documents/"))
        self.assertTrue(
            PortalAuditLog.objects.filter(
                action="create",
                target_model="news.news",
                target_object_id=str(news.pk),
            ).exists()
        )

    def test_news_import_requires_overwrite_for_existing_article(self):
        News.objects.create(
            title="BÀI TEST IMPORT",
            slug="bai-test-import",
            content="old content",
            category=self.category,
        )

        response = self.client.post(
            reverse("portal:news_import"),
            data={
                "source_file": self._make_docx_upload(
                    "BÀI TEST IMPORT",
                    ["Nội dung mới."],
                ),
                "category": self.category.pk,
            },
            HTTP_HOST="localhost",
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Bài viết đã tồn tại")
        self.assertEqual(News.all_objects.filter(slug="bai-test-import").count(), 1)

    def test_news_import_can_overwrite_existing_article(self):
        existing = News.objects.create(
            title="BÀI TEST IMPORT",
            slug="bai-test-import",
            content="old content",
            excerpt="old",
            category=self.category,
        )

        response = self.client.post(
            reverse("portal:news_import"),
            data={
                "source_file": self._make_docx_upload(
                    "BÀI TEST IMPORT",
                    ["Nội dung đã được cập nhật."],
                ),
                "category": self.category.pk,
                "overwrite_existing": "on",
            },
            HTTP_HOST="localhost",
        )

        self.assertEqual(response.status_code, 302)
        existing.refresh_from_db()
        self.assertIn("Nội dung đã được cập nhật", existing.content)
        self.assertTrue(existing.source_document.name.startswith("news/source-documents/"))
