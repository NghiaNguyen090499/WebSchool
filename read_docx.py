import docx
doc = docx.Document("d:/NGHIA/WebsiteSchool/Tổng quan chương trình GD MIS 2026-2027.docx")
with open("d:/NGHIA/WebsiteSchool/parsed_docx.txt", "w", encoding="utf-8") as f:
    for p in doc.paragraphs:
        f.write(p.text + "\n")
    for table in doc.tables:
        for row in table.rows:
            f.write("\t".join([cell.text for cell in row.cells]) + "\n")
