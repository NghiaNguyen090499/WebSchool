from __future__ import annotations

import argparse
import hashlib
import os
import re
import shutil
import sys
import unicodedata
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from html import escape
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

os.environ["DEBUG"] = "False"
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "school_website.settings")

import django

django.setup()

from django.conf import settings

from news.models import Category, News


SOURCE_DIR = REPO_ROOT / "Bài viết Tháng 4"
EXTRA_IMAGES_DIR = SOURCE_DIR / "Ảnh HSG đăng thêm vào gương mặt"
MEDIA_ROOT = Path(settings.MEDIA_ROOT)
MEDIA_URL = settings.MEDIA_URL.rstrip("/")
LOCAL_TZ = timezone(timedelta(hours=7))
NS_A = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
URL_RE = re.compile(r"(https?://[^\s<]+)")


@dataclass(frozen=True)
class ArticleConfig:
    category_name: str
    extra_images_dir: str | Path | None = None


ARTICLE_CONFIGS = [
    (
        "mis-ghi-dau-an-tai-ky-thi-hoc-gioi",
        ArticleConfig(category_name="Gương mặt Misers"),
    ),
    (
        "mis-innovation-day-the-new-era-2026",
        ArticleConfig(category_name="Sự kiện nhà trường"),
    ),
    (
        "mis-ky-ket-hop-tac-chien-luoc-voi-ismart-va-edulive",
        ArticleConfig(category_name="Hợp tác quốc tế"),
    ),
    (
        "mis-vinh-danh-hoc-sinh-le-khanh-hoa",
        ArticleConfig(
            category_name="Gương mặt Misers",
            extra_images_dir="Ảnh HSG đăng thêm vào gương mặt",
        ),
    ),
    (
        "misers-toa-sang-tai-cuoc-thi-hung-bien-tieng-anh-soundbites-2026",
        ArticleConfig(category_name="Cuộc thi & Học thuật"),
    ),
    (
        "mis-innovation-day-2026-ban",
        ArticleConfig(category_name="Sự kiện nhà trường"),
    ),
]


def get_article_config(slug: str) -> ArticleConfig:
    for prefix, config in ARTICLE_CONFIGS:
        if slug.startswith(prefix):
            return config
    raise KeyError(f"No import config for article slug: {slug}")


def iter_block_items(document: Document) -> Iterable[Paragraph | Table]:
    parent = document.element.body
    for child in parent.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def clean_text(value: str) -> str:
    value = value.replace("\xa0", " ").replace("\u2028", " ").replace("\u2029", " ")
    value = value.replace("\u200b", "")
    value = value.replace("\r", "\n")
    lines = [re.sub(r"\s+", " ", line).strip() for line in value.split("\n")]
    return "\n".join(line for line in lines if line)


def make_slug(value: str) -> str:
    value = clean_text(value)
    value = unicodedata.normalize("NFKC", value)
    value = value.replace("Đ", "D").replace("đ", "d")
    value = unicodedata.normalize("NFD", value)
    value = "".join(char for char in value if not unicodedata.combining(char))
    value = re.sub(r"[^A-Za-z0-9]+", "-", value)
    value = re.sub(r"-{2,}", "-", value)
    return value.strip("-").lower()


def paragraph_lines(paragraph: Paragraph) -> list[str]:
    text = clean_text(paragraph.text)
    if not text:
        return []
    return [line for line in text.split("\n") if line]


def paragraph_images(paragraph: Paragraph, document: Document) -> list[tuple[str, bytes]]:
    images: list[tuple[str, bytes]] = []
    seen_ids: set[str] = set()
    for run in paragraph.runs:
        for blip in run.element.findall(".//a:blip", namespaces=NS_A):
            embed_id = blip.get(qn("r:embed"))
            if not embed_id or embed_id in seen_ids:
                continue
            part = document.part.related_parts.get(embed_id)
            if not part:
                continue
            content_type = getattr(part, "content_type", "")
            if not content_type.startswith("image/"):
                continue
            ext = content_type.split("/")[-1].replace("jpeg", "jpg")
            images.append((ext, part.blob))
            seen_ids.add(embed_id)
    return images


def save_binary(relative_path: str, payload: bytes, dry_run: bool) -> str:
    normalized = relative_path.replace("\\", "/")
    absolute_path = MEDIA_ROOT / Path(normalized)
    if not dry_run:
        absolute_path.parent.mkdir(parents=True, exist_ok=True)
        absolute_path.write_bytes(payload)
    return normalized


def copy_extra_images(source_dir: Path, target_dir: str, start_index: int, dry_run: bool) -> list[str]:
    paths: list[str] = []
    for offset, image_path in enumerate(sorted(source_dir.glob("*"))):
        if not image_path.is_file():
            continue
        suffix = image_path.suffix.lower()
        if suffix not in {".jpg", ".jpeg", ".png", ".webp"}:
            continue
        relative_path = f"{target_dir}/image_{start_index + offset:02d}{suffix}"
        normalized = relative_path.replace("\\", "/")
        absolute_target = MEDIA_ROOT / Path(normalized)
        if not dry_run:
            absolute_target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(image_path, absolute_target)
        paths.append(normalized)
    return paths


def html_url(relative_media_path: str) -> str:
    return f"{MEDIA_URL}/{relative_media_path}".replace("//", "/").replace(":/", "://")


def looks_like_short_item(text: str) -> bool:
    if len(text) > 120:
        return False
    if re.search(r"\bGiải\b", text, flags=re.IGNORECASE):
        return True
    if re.search(r" – ", text):
        return True
    if ":" in text and len(text) < 100:
        return True
    return text.endswith(".") and len(text) < 90


def render_text_blocks(lines: list[str]) -> list[str]:
    blocks: list[str] = []
    index = 0
    while index < len(lines):
        current = lines[index]
        next_lines = lines[index + 1 :]
        if current.endswith(":"):
            list_items: list[str] = []
            for item in next_lines:
                if looks_like_short_item(item):
                    list_items.append(item)
                else:
                    break
            if list_items:
                blocks.append(f"<p><strong>{linkify_text(current)}</strong></p>")
                blocks.append(
                    "<ul>"
                    + "".join(f"<li>{linkify_text(item)}</li>" for item in list_items)
                    + "</ul>"
                )
                index += 1 + len(list_items)
                continue
        blocks.append(f"<p>{linkify_text(current)}</p>")
        index += 1
    return blocks


def linkify_text(text: str) -> str:
    escaped = escape(text, quote=False)

    def replace(match: re.Match[str]) -> str:
        url = match.group(1)
        safe_url = escape(url, quote=True)
        label = escape(url, quote=False)
        return f'<a href="{safe_url}" target="_blank" rel="noopener">{label}</a>'

    return URL_RE.sub(replace, escaped)


def build_article_payload(docx_path: Path, config: ArticleConfig, dry_run: bool) -> dict:
    document = Document(str(docx_path))
    raw_title = next((clean_text(p.text) for p in document.paragraphs if clean_text(p.text)), docx_path.stem)
    title = raw_title
    slug = make_slug(title)
    media_dir = f"news/imported/2026-04/{slug[:80]}"

    body_lines: list[str] = []
    image_paths: list[str] = []
    seen_image_hashes: set[str] = set()
    skip_first_title = True
    image_counter = 1

    for block in iter_block_items(document):
        if isinstance(block, Table):
            continue

        lines = paragraph_lines(block)
        images = paragraph_images(block, document)

        if lines:
            if skip_first_title and lines[0] == title:
                lines = lines[1:]
                skip_first_title = False
            else:
                skip_first_title = False

        if lines:
            body_lines.extend(lines)

        for ext, payload in images:
            payload_hash = hashlib.sha1(payload).hexdigest()
            if payload_hash in seen_image_hashes:
                continue
            image_path = save_binary(
                f"{media_dir}/image_{image_counter:02d}.{ext}",
                payload,
                dry_run=dry_run,
            )
            image_paths.append(image_path)
            seen_image_hashes.add(payload_hash)
            image_counter += 1

    extra_images_dir: Path | None = None
    if isinstance(config.extra_images_dir, Path):
        extra_images_dir = config.extra_images_dir
    elif isinstance(config.extra_images_dir, str):
        extra_images_dir = docx_path.parent / config.extra_images_dir

    if extra_images_dir and extra_images_dir.exists():
        image_paths.extend(
            copy_extra_images(
                source_dir=extra_images_dir,
                target_dir=media_dir,
                start_index=image_counter,
                dry_run=dry_run,
            )
        )

    html_blocks = render_text_blocks(body_lines)
    if image_paths:
        image_blocks = [
            f'<figure><img src="{html_url(path)}" alt="{escape(title, quote=True)}"></figure>'
            for path in image_paths
        ]
        html_blocks.extend(image_blocks)

    excerpt_source = " ".join(body_lines[:3]).strip()
    excerpt = excerpt_source[:297].rstrip()
    if len(excerpt_source) > len(excerpt):
        excerpt = f"{excerpt}..."

    created_at = datetime.fromtimestamp(docx_path.stat().st_mtime, tz=LOCAL_TZ)
    category = Category.objects.get(name=config.category_name)

    return {
        "title": title,
        "slug": slug,
        "category": category,
        "content": "\n".join(html_blocks),
        "excerpt": excerpt,
        "thumbnail": image_paths[0] if image_paths else None,
        "created_at": created_at,
        "image_count": len(image_paths),
        "source_path": docx_path,
    }


def import_articles(source_dir: Path, dry_run: bool) -> None:
    processed = 0
    for docx_path in sorted(source_dir.glob("*.docx")):
        document = Document(str(docx_path))
        title = next((clean_text(p.text) for p in document.paragraphs if clean_text(p.text)), docx_path.stem)
        slug = make_slug(title)
        config = get_article_config(slug)

        payload = build_article_payload(docx_path, config=config, dry_run=dry_run)
        article = News.all_objects.filter(title=payload["title"]).first()
        if not article:
            article = News.all_objects.filter(slug=payload["slug"]).first()
        action = "update" if article else "create"

        print(f"[{action.upper()}] {payload['title']}")
        print(f"  category: {payload['category'].name}")
        print(f"  created_at: {payload['created_at'].isoformat()}")
        print(f"  images: {payload['image_count']}")

        if dry_run:
            processed += 1
            continue

        if article:
            article.title = payload["title"]
            article.content = payload["content"]
            article.excerpt = payload["excerpt"]
            article.category = payload["category"]
            article.is_featured = False
            article.is_archived = False
            article.slug = payload["slug"]
        else:
            article = News(
                title=payload["title"],
                slug=payload["slug"],
                content=payload["content"],
                excerpt=payload["excerpt"],
                category=payload["category"],
                is_featured=False,
                is_archived=False,
            )

        if payload["thumbnail"]:
            article.thumbnail.name = payload["thumbnail"]

        article.save()
        News.all_objects.filter(pk=article.pk).update(created_at=payload["created_at"])
        processed += 1

    print(f"Processed {processed} article(s).")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Import April news articles from docx files.")
    parser.add_argument(
        "--source-dir",
        type=Path,
        default=SOURCE_DIR,
        help="Directory containing the April docx articles.",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Parse the articles and print the import plan without writing files or DB rows.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    import_articles(source_dir=args.source_dir, dry_run=args.dry_run)


if __name__ == "__main__":
    main()
