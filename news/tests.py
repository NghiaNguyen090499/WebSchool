import shutil
import tempfile
from io import BytesIO

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from docx import Document

from .models import Category, News
from .word_export import build_news_docx_bytes, export_news_to_source_document


class NewsRichTextSanitizerTests(TestCase):
    def setUp(self):
        super().setUp()
        self.temp_media = tempfile.mkdtemp()
        self.addCleanup(shutil.rmtree, self.temp_media, ignore_errors=True)
        self.override_media = override_settings(MEDIA_ROOT=self.temp_media)
        self.override_media.enable()
        self.addCleanup(self.override_media.disable)

    def test_news_detail_sanitizes_html(self):
        category = Category.objects.create(name="General")
        news = News.objects.create(
            title="Hello",
            content="<script>alert(1)</script><p>Hello</p>",
            category=category,
        )

        response = self.client.get(news.get_absolute_url(), HTTP_HOST="localhost")
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "<script>alert(1)</script>")
        self.assertContains(response, "<p>Hello</p>", html=True)

    def test_news_detail_shows_download_link_for_source_document(self):
        category = Category.objects.create(name="Downloads")
        source_document = SimpleUploadedFile(
            "source.docx",
            b"PK\x03\x04word-test",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        news = News.objects.create(
            title="Word Source",
            content="<p>Hello</p>",
            category=category,
            source_document=source_document,
        )

        response = self.client.get(news.get_absolute_url(), HTTP_HOST="localhost")
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Tải file Word gốc")
        self.assertContains(response, news.source_document.url)


class NewsWordExportTests(TestCase):
    def setUp(self):
        super().setUp()
        self.temp_media = tempfile.mkdtemp()
        self.addCleanup(shutil.rmtree, self.temp_media, ignore_errors=True)
        self.override_media = override_settings(MEDIA_ROOT=self.temp_media)
        self.override_media.enable()
        self.addCleanup(self.override_media.disable)

    def test_build_news_docx_bytes_preserves_text_links_and_images(self):
        category = Category.objects.create(name="Export")
        news = News.objects.create(
            title="Exported News",
            excerpt="Summary line",
            content="""
                <div class="jeg_share_button">Share</div>
                <div class="content-inner">
                    <h2>Section title</h2>
                    <p>Hello <a href="https://example.com">example</a></p>
                    <ul><li>First item</li><li>Second item</li></ul>
                    <figure><img src="/media/news/content/demo.jpg" /></figure>
                </div>
            """,
            category=category,
        )

        payload = build_news_docx_bytes(news, site_base_url="https://misvn.edu.vn")
        document = Document(BytesIO(payload))
        full_text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

        self.assertIn("Exported News", full_text)
        self.assertIn("Summary line", full_text)
        self.assertIn("Section title", full_text)
        self.assertIn("Hello example (https://example.com)", full_text)
        self.assertIn("First item", full_text)
        self.assertIn("https://misvn.edu.vn/media/news/content/demo.jpg", full_text)

    def test_export_news_to_source_document_generates_docx_for_missing_source(self):
        category = Category.objects.create(name="Generated")
        news = News.objects.create(
            title="Generated Word",
            content="<p>Hello world from website.</p>",
            category=category,
        )

        result = export_news_to_source_document(news)
        news.refresh_from_db()

        self.assertEqual(result.action, "create")
        self.assertTrue(news.source_document.name.startswith("news/source-documents/generated/"))

        document = Document(news.source_document.path)
        full_text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        self.assertIn("Generated Word", full_text)
        self.assertIn("Hello world from website.", full_text)

    def test_export_news_to_source_document_keeps_existing_file_without_overwrite(self):
        category = Category.objects.create(name="Existing")
        news = News.objects.create(
            title="Original Word",
            content="<p>Old</p>",
            category=category,
            source_document=SimpleUploadedFile(
                "original.docx",
                b"PK\x03\x04existing",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )
        original_name = news.source_document.name

        result = export_news_to_source_document(news, overwrite=False)
        news.refresh_from_db()

        self.assertEqual(result.action, "skip")
        self.assertEqual(news.source_document.name, original_name)
