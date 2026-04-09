from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import re
from typing import Iterable
from urllib.parse import urljoin

from django.core.files.base import ContentFile
from django.db.models import Q, QuerySet
from django.utils import timezone
from django.utils.html import strip_tags
from django.utils.text import slugify
from docx import Document

from .models import News

try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except ImportError:  # pragma: no cover - fallback only used if bs4 is unavailable.
    BeautifulSoup = None
    NavigableString = str  # type: ignore[assignment]
    Tag = object  # type: ignore[assignment]


HEADING_LEVELS = {
    "h1": 1,
    "h2": 2,
    "h3": 3,
    "h4": 4,
    "h5": 4,
    "h6": 4,
}
CONTAINER_TAGS = {"article", "section", "main", "div"}
BLOCK_TAGS = set(HEADING_LEVELS) | {
    "p",
    "ul",
    "ol",
    "li",
    "blockquote",
    "table",
    "figure",
    "img",
} | CONTAINER_TAGS
NOISE_SELECTORS = (
    ".jeg_share_button",
    ".jeg_share_float_container",
    ".jeg_sharelist",
    ".share-secondary",
    ".share-float",
    ".sharedaddy",
    ".jp-relatedposts",
)


@dataclass(frozen=True)
class NewsWordExportResult:
    action: str
    news_id: int
    slug: str
    document_name: str | None = None


def build_news_docx_bytes(news: News, site_base_url: str | None = None) -> bytes:
    document = Document()

    document.add_heading(news.title, level=0)

    metadata_lines = [
        f"Slug: {news.slug}",
        f"Chuyen muc: {news.category.name if news.category else 'Chua phan loai'}",
        f"Ngay dang: {_format_datetime(news.created_at)}",
        f"Duong dan: {_build_url(news.get_absolute_url(), site_base_url)}",
        "Nguon tai lieu: Tep Word duoc tao tu dong tu noi dung tren website.",
    ]
    for line in metadata_lines:
        document.add_paragraph(line)

    if news.excerpt:
        document.add_heading("Tom tat", level=1)
        document.add_paragraph(news.excerpt)

    document.add_heading("Noi dung", level=1)
    image_sources = append_news_content_to_document(
        document=document,
        html=news.content,
        site_base_url=site_base_url,
    )

    if image_sources:
        document.add_heading("Hinh anh tren website", level=1)
        for src in image_sources:
            document.add_paragraph(src, style="List Bullet")

    payload = BytesIO()
    document.save(payload)
    return payload.getvalue()


def append_news_content_to_document(
    document: Document,
    html: str,
    site_base_url: str | None = None,
) -> list[str]:
    if not html:
        return []

    if BeautifulSoup is None:
        plain_text = _fallback_plain_text(html)
        for paragraph in plain_text.splitlines():
            text = _normalize_text(paragraph)
            if text:
                document.add_paragraph(text)
        return []

    soup = BeautifulSoup(html, "html.parser")
    _strip_noise(soup)
    root = soup.select_one(".content-inner") or soup.select_one(".entry-content") or soup
    image_sources: list[str] = []
    seen_images: set[str] = set()

    for child in root.children:
        _append_node(
            document=document,
            node=child,
            image_sources=image_sources,
            seen_images=seen_images,
            site_base_url=site_base_url,
        )

    return image_sources


def export_news_to_source_document(
    news: News,
    *,
    overwrite: bool = False,
    site_base_url: str | None = None,
) -> NewsWordExportResult:
    if news.source_document and not overwrite:
        return NewsWordExportResult(
            action="skip",
            news_id=news.pk,
            slug=news.slug,
            document_name=news.source_document.name,
        )

    if news.source_document and overwrite:
        news.source_document.delete(save=False)

    filename = build_generated_document_name(news)
    payload = build_news_docx_bytes(news, site_base_url=site_base_url)
    news.source_document.save(filename, ContentFile(payload), save=False)
    News.all_objects.filter(pk=news.pk).update(source_document=news.source_document.name)
    news.refresh_from_db(fields=["source_document"])

    return NewsWordExportResult(
        action="replace" if overwrite else "create",
        news_id=news.pk,
        slug=news.slug,
        document_name=news.source_document.name,
    )


def export_news_queryset_to_source_documents(
    queryset: QuerySet[News] | Iterable[News],
    *,
    overwrite: bool = False,
    site_base_url: str | None = None,
) -> list[NewsWordExportResult]:
    results: list[NewsWordExportResult] = []
    iterable = queryset.iterator(chunk_size=100) if isinstance(queryset, QuerySet) else queryset
    for news in iterable:
        results.append(
            export_news_to_source_document(
                news,
                overwrite=overwrite,
                site_base_url=site_base_url,
            )
        )
    return results


def build_missing_source_document_queryset() -> QuerySet[News]:
    return News.all_objects.filter(Q(source_document__isnull=True) | Q(source_document="")).order_by("pk")


def build_generated_document_name(news: News) -> str:
    created_at = timezone.localtime(news.created_at) if timezone.is_aware(news.created_at) else news.created_at
    month_folder = created_at.strftime("%Y-%m") if created_at else "undated"
    slug = news.slug or slugify(news.title) or f"news-{news.pk}"
    return f"generated/{month_folder}/{news.pk}-{slug}.docx"


def _append_node(
    document: Document,
    node,
    image_sources: list[str],
    seen_images: set[str],
    site_base_url: str | None,
) -> None:
    if BeautifulSoup is None:
        return

    if isinstance(node, NavigableString):
        text = _normalize_text(str(node))
        if text:
            document.add_paragraph(text)
        return

    if not isinstance(node, Tag):
        return

    if node.name in {"script", "style", "noscript", "svg"}:
        return

    if node.name == "img":
        _remember_image(node, image_sources, seen_images, site_base_url)
        return

    if node.name == "figure":
        for image in node.find_all("img", recursive=False):
            _remember_image(image, image_sources, seen_images, site_base_url)
        caption = _extract_text_with_links(node.find("figcaption"))
        if caption:
            document.add_paragraph(caption)
        return

    if node.name in HEADING_LEVELS:
        text = _extract_text_with_links(node)
        if text:
            document.add_heading(text, level=HEADING_LEVELS[node.name])
        return

    if node.name in {"ul", "ol"}:
        style_name = "List Bullet" if node.name == "ul" else "List Number"
        for item in node.find_all("li", recursive=False):
            text = _extract_text_with_links(item)
            if text:
                document.add_paragraph(text, style=style_name)
        return

    if node.name == "table":
        for row in node.find_all("tr"):
            cells = [_normalize_text(cell.get_text(" ", strip=True)) for cell in row.find_all(["th", "td"])]
            cells = [cell for cell in cells if cell]
            if cells:
                document.add_paragraph(" | ".join(cells))
        return

    if node.name == "blockquote":
        text = _extract_text_with_links(node)
        if text:
            paragraph = document.add_paragraph(text)
            for run in paragraph.runs:
                run.italic = True
        return

    if node.name == "p":
        for image in node.find_all("img", recursive=False):
            _remember_image(image, image_sources, seen_images, site_base_url)
        text = _extract_text_with_links(node)
        if text:
            document.add_paragraph(text)
        return

    if node.name in CONTAINER_TAGS:
        if _has_direct_block_children(node):
            for child in node.children:
                _append_node(document, child, image_sources, seen_images, site_base_url)
            return
        text = _extract_text_with_links(node)
        if text:
            document.add_paragraph(text)
        return

    if list(node.children):
        for child in node.children:
            _append_node(document, child, image_sources, seen_images, site_base_url)
        return

    text = _extract_text_with_links(node)
    if text:
        document.add_paragraph(text)


def _strip_noise(soup) -> None:
    for tag_name in ("script", "style", "noscript", "svg"):
        for node in soup.find_all(tag_name):
            node.decompose()
    for selector in NOISE_SELECTORS:
        for node in soup.select(selector):
            node.decompose()


def _has_direct_block_children(tag) -> bool:
    if BeautifulSoup is None:
        return False
    for child in tag.children:
        if isinstance(child, Tag) and child.name in BLOCK_TAGS:
            return True
    return False


def _extract_text_with_links(tag) -> str:
    if not tag:
        return ""
    if BeautifulSoup is None:
        return _normalize_text(strip_tags(str(tag)))

    clone = BeautifulSoup(str(tag), "html.parser")
    for node in clone.find_all(["script", "style", "noscript", "svg", "img"]):
        node.decompose()
    for anchor in clone.find_all("a"):
        label = _normalize_text(anchor.get_text(" ", strip=True))
        href = _normalize_text(anchor.get("href", ""))
        if href and label and href not in label:
            replacement = f"{label} ({href})"
        else:
            replacement = label or href
        anchor.replace_with(replacement)
    return _normalize_text(clone.get_text(" ", strip=True))


def _remember_image(tag, image_sources: list[str], seen_images: set[str], site_base_url: str | None) -> None:
    src = _normalize_text(tag.get("src", ""))
    if not src or "emoji.php" in src:
        return
    image_url = _build_url(src, site_base_url)
    if image_url in seen_images:
        return
    seen_images.add(image_url)
    image_sources.append(image_url)


def _fallback_plain_text(html: str) -> str:
    html = re.sub(r"(?i)<\s*br\s*/?\s*>", "\n", html)
    html = re.sub(r"(?i)</\s*(p|div|li|h1|h2|h3|h4|h5|h6|tr)\s*>", "\n", html)
    return strip_tags(html)


def _normalize_text(value: str) -> str:
    value = (value or "").replace("\xa0", " ")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _format_datetime(value) -> str:
    if not value:
        return ""
    if timezone.is_aware(value):
        value = timezone.localtime(value)
    return value.strftime("%d/%m/%Y %H:%M")


def _build_url(path_or_url: str, site_base_url: str | None) -> str:
    if not path_or_url:
        return ""
    if path_or_url.startswith(("http://", "https://")):
        return path_or_url
    if site_base_url:
        return urljoin(site_base_url.rstrip("/") + "/", path_or_url.lstrip("/"))
    return path_or_url
